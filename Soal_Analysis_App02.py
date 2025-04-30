import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.io as pio
import matplotlib.pyplot as plt
from docx.shared import Inches
import tempfile
from io import BytesIO
from docx import Document

if 'hasil' in st.session_state:
    df_hasil = pd.DataFrame(st.session_state.hasil, columns=["Soal", "Tipe", "Nilai", "Keterangan"])
# -------------------------
# Konfigurasi Halaman
# -------------------------
st.set_page_config(page_title="Analisis Soal Interaktif", layout="wide")

# -------------------------
# Sidebar Navigasi
# -------------------------
with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/commons/thumb/1/17/Google-flutter-logo.png/768px-Google-flutter-logo.png", width=140)
    st.markdown("### 📊 Analisis Soal")
    menu = st.radio("Navigasi", ["🏠 Beranda", "📤 Unggah Data", "📈 Analisis", "📥 Unduh Hasil", "ℹ️ Tentang"])

# -------------------------
# Header
# -------------------------
st.markdown("""
    <div style='background-color:#4B8BBE; padding:15px; border-radius:10px; text-align:center'>
        <h2 style='color:white;'>📘 Aplikasi Analisis Soal Interaktif</h2>
    </div>
""", unsafe_allow_html=True)

# -------------------------
# Halaman Beranda
# -------------------------
if menu == "🏠 Beranda":
    st.subheader("Selamat Datang di Aplikasi Analisis Soal!")
    st.markdown("""
        Selamat datang di **Aplikasi Analisis Soal**! Aplikasi ini dirancang untuk membantu Anda menganalisis hasil ujian atau kuis berdasarkan jenis soal yang diberikan.
        
        Aplikasi ini membantu untuk mendapatkan insight yang lebih baik mengenai performa siswa dalam ujian. Silakan unggah file jawaban dan mulailah analisis!
    """)


# -------------------------
# Halaman Unggah Data
# -------------------------
# -------------------------
# Halaman Unggah Data
# -------------------------
elif menu == "📤 Unggah Data":
    st.subheader("📤 Unggah File Jawaban")
    uploaded_file = st.file_uploader("Unggah file Excel atau CSV", type=["xlsx", "csv"])

    # Input Nama Pengguna dan Mata Pelajaran
    nama_user = st.text_input("Masukkan Nama Pengguna")
    mapel = st.text_input("Masukkan Mata Pelajaran")

    if uploaded_file:
        df = pd.read_csv(uploaded_file) if uploaded_file.name.endswith('.csv') else pd.read_excel(uploaded_file)
        st.success("✅ File berhasil diunggah!")
        st.dataframe(df)

        st.markdown("### Kategori Soal")
        soal_cols = df.columns[1:]

        if 'kategori' not in st.session_state:
            st.session_state.kategori = {}

        for col in soal_cols:
            kategori = st.selectbox(f"{col}", ["PG", "Isian", "Esai", "Belum Tersedia", "Tambah Manual"], key=col)
            if kategori == "Tambah Manual":
                manual_input = st.text_input(f"Masukkan kategori untuk {col}", key=f"input_{col}")
                st.session_state.kategori[col] = manual_input or "Belum Tersedia"
            else:
                st.session_state.kategori[col] = kategori

        # Menyimpan nama pengguna dan mapel ke session_state
        st.session_state.nama_user = nama_user
        st.session_state.mapel = mapel

        if st.button("✅ Simpan & Lanjutkan ke Analisis"):
            st.session_state.df = df
            st.success("Kategori berhasil disimpan. Lanjut ke menu **Analisis**.")


# -------------------------
# Halaman Analisis
# -------------------------
# -------------------------
# -------------------------
# Halaman Analisis
# -------------------------
elif menu == "📈 Analisis":
    if 'df' not in st.session_state or 'kategori' not in st.session_state:
        st.warning("⚠️ Silakan unggah data terlebih dahulu.")
    else:
        df = st.session_state.df
        kategori = st.session_state.kategori
        hasil = []

        # Menampilkan Nama Pengguna dan Mata Pelajaran
        nama_user = st.session_state.get('nama_user', 'Tidak ada nama pengguna')
        mapel = st.session_state.get('mapel', 'Tidak ada mata pelajaran')

        st.subheader(f"📊 Hasil Analisis Soal untuk {nama_user} - Mata Pelajaran: {mapel}")

        # Analisis Soal
        for col, tipe in kategori.items():
            jawaban_benar = df[col][0]
            jawaban_siswa = df[col][1:]

            if tipe == "PG":
                total = len(jawaban_siswa)
                benar = (jawaban_siswa == jawaban_benar).sum()
                p = benar / total if total else 0
                tingkat = "Mudah" if p >= 0.7 else "Sedang" if p >= 0.3 else "Sulit"
                hasil.append([col, tipe, f"{p*100:.2f}%", tingkat])

            elif tipe == "Isian":
                total = len(jawaban_siswa)
                benar = (jawaban_siswa.astype(str).str.lower().str.strip() == str(jawaban_benar).lower().strip()).sum()
                p = benar / total if total else 0
                tingkat = "Mudah" if p >= 0.7 else "Sedang" if p >= 0.3 else "Sulit"
                hasil.append([col, tipe, f"{p*100:.2f}%", tingkat])

            elif tipe == "Esai":
                try:
                    max_skor = float(jawaban_benar)
                    skor = pd.to_numeric(jawaban_siswa, errors='coerce')
                    mean_skor = skor.mean()
                    std_dev = skor.std(ddof=0)
                    tingkat = "Mudah" if mean_skor >= 4.0 else "Sedang" if mean_skor >= 2.5 else "Sulit"
                    hasil.append([col, tipe, f"{mean_skor:.2f} / {max_skor}", f"STD: {std_dev:.2f} - {tingkat}"])
                except:
                    hasil.append([col, tipe, "Invalid", "Format salah"])

            else:
                hasil.append([col, tipe, "N/A", "Belum dikategorikan"])

        hasil_df = pd.DataFrame(hasil, columns=["Soal", "Tipe", "Nilai", "Keterangan"])
        st.dataframe(hasil_df)
        st.session_state.hasil = hasil

        # Grafik Visualisasi
        if st.checkbox("📈 Tampilkan Grafik Visualisasi", key="grafik_visual"):
            pg_data = hasil_df[hasil_df['Tipe'] == 'PG']
            if not pg_data.empty:
                pg_data['Nilai (%)'] = pg_data['Nilai'].str.replace('%', '').astype(float)
                fig_pg = px.bar(pg_data, x="Soal", y="Nilai (%)", color="Keterangan", title=f"Distribusi Soal PG - {mapel}")
                st.plotly_chart(fig_pg, use_container_width=True)

            isian_data = hasil_df[hasil_df['Tipe'] == 'Isian']
            if not isian_data.empty:
                isian_data['Nilai (%)'] = isian_data['Nilai'].str.replace('%', '').astype(float)
                fig_isian = px.pie(isian_data, values="Nilai (%)", names="Soal", title=f"Distribusi Isian Singkat - {mapel}")
                st.plotly_chart(fig_isian, use_container_width=True)

            esai_data = hasil_df[hasil_df['Tipe'] == 'Esai']
            if not esai_data.empty:
                try:
                    nilai_angka = esai_data["Nilai"].str.extract(r"(\d+\.\d+)")[0].astype(float)
                    fig_esai = px.scatter(x=esai_data["Soal"], y=nilai_angka, color=esai_data["Keterangan"], title=f"Rata-Rata Nilai Esai - {mapel}")
                    st.plotly_chart(fig_esai, use_container_width=True)
                except:
                    st.warning("⚠️ Gagal memvisualisasikan nilai esai.")

        # Histogram Esai
        if st.checkbox("📊 Tampilkan Histogram Nilai Esai", key="histogram_esai"):
            for col, tipe in kategori.items():
                if tipe == "Esai":
                    try:
                        nilai_angka = pd.to_numeric(df[col][1:], errors='coerce').dropna()
                        if not nilai_angka.empty:
                            fig_hist = plt.figure()
                            plt.hist(nilai_angka, bins=10, color='skyblue', edgecolor='black')
                            plt.title(f"Distribusi Nilai {col}")
                            plt.xlabel("Nilai")
                            plt.ylabel("Frekuensi")
                            st.pyplot(fig_hist)
                        else:
                            st.info(f"ℹ️ Tidak ada data numerik untuk kolom {col}.")
                    except:
                        st.warning(f"⚠️ Gagal memvisualisasikan histogram nilai untuk {col}.")



# -------------------------
elif menu == "📥 Unduh Hasil":
    if 'hasil' in st.session_state:
        import pandas as pd
        from io import BytesIO
        from docx import Document
        from docx.shared import Inches
        import matplotlib.pyplot as plt
        import streamlit as st

        # Ambil data hasil
        df_pg = pd.DataFrame(st.session_state.hasil, columns=["Soal", "Tipe", "Nilai", "Keterangan"])
        df_esai = pd.DataFrame(st.session_state.get("hasil_esai", []), columns=["Soal", "Jawaban", "Skor", "Komentar"])
        df_isian = pd.DataFrame(st.session_state.get("hasil_isian", []), columns=["Soal", "Jawaban", "Skor", "Komentar"])

        # ========== UNDUH EXCEL ==========
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
            df_pg.to_excel(writer, index=False, sheet_name="Pilihan Ganda")
            if not df_esai.empty:
                df_esai.to_excel(writer, index=False, sheet_name="Esai")
            if not df_isian.empty:
                df_isian.to_excel(writer, index=False, sheet_name="Isian")

        st.download_button(
            "📥 Unduh sebagai Excel",
            data=excel_buffer.getvalue(),
            file_name="hasil_analisis.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # ========== UNDUH WORD ==========
        def convert_to_float(x):
            if isinstance(x, str):
                x = x.strip()
                if '%' in x:
                    return float(x.replace('%', ''))
                elif '/' in x:
                    try:
                        num, denom = x.split('/')
                        return float(num.strip()) / float(denom.strip()) * 100  # jadikan persen
                    except:
                        return None
                else:
                    try:
                        return float(x)
                    except:
                        return None
            return x

        nilai_pg = df_pg["Nilai"].apply(convert_to_float)
        skor_esai = df_esai["Skor"].apply(convert_to_float) if not df_esai.empty else []
        skor_isian = df_isian["Skor"].apply(convert_to_float) if not df_isian.empty else []

        doc = Document()
        doc.add_heading("📊 Hasil Analisis Lengkap", 0)

        # Info user
        nama_user = st.session_state.get('nama_user', 'Tidak ada nama pengguna')
        mapel = st.session_state.get('mapel', 'Tidak ada mata pelajaran')
        doc.add_paragraph(f"👤 Nama Pengguna: {nama_user}")
        doc.add_paragraph(f"📚 Mata Pelajaran: {mapel}")
        doc.add_paragraph("")

        # ======= Bagian Pilihan Ganda =======
        doc.add_heading("📋 Analisis Pilihan Ganda", level=1)
        table_pg = doc.add_table(rows=1, cols=len(df_pg.columns))
        table_pg.style = 'Table Grid'
        for i, col in enumerate(df_pg.columns):
            table_pg.rows[0].cells[i].text = col
        for row in df_pg.itertuples(index=False):
            cells = table_pg.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

        # Grafik PG
        if not df_pg.empty:
            fig, ax = plt.subplots(figsize=(6, 4))
            bars = ax.bar(df_pg["Soal"], nilai_pg, color="skyblue")
            ax.set_title("Grafik Nilai Pilihan Ganda")
            ax.set_xlabel("Soal")
            ax.set_ylabel("Nilai (%)")
            plt.xticks(rotation=45)
            for bar, value in zip(bars, nilai_pg):
                if pd.notnull(value):
                    ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), f"{value:.1f}",
                            ha='center', va='bottom', fontsize=8)  # Perbaikan di sini
            pg_stream = BytesIO()
            plt.tight_layout()
            plt.savefig(pg_stream, format='png')
            plt.close(fig)
            pg_stream.seek(0)
            doc.add_paragraph("\n📈 Grafik Nilai Pilihan Ganda:")
            doc.add_picture(pg_stream, width=Inches(6 ))

        # ======= Bagian Esai =======
        if not df_esai.empty:
            doc.add_page_break()
            doc.add_heading("📝 Analisis Esai", level=1)
            table_esai = doc.add_table(rows=1, cols=len(df_esai.columns))
            table_esai.style = 'Table Grid'
            for i, col in enumerate(df_esai.columns):
                table_esai.rows[0].cells[i].text = col
            for row in df_esai.itertuples(index=False):
                cells = table_esai.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)

            # Grafik Esai jika skor numerik tersedia
            if skor_esai and any(pd.notnull(skor_esai)):
                fig2, ax2 = plt.subplots(figsize=(6, 4))
                bars2 = ax2.bar(df_esai["Soal"], skor_esai, color="lightcoral")
                ax2.set_title("Grafik Skor Esai")
                ax2.set_xlabel("Soal Esai")
                ax2.set_ylabel("Skor (%)")
                plt.xticks(rotation=45)
                for bar, value in zip(bars2, skor_esai):
                    if pd.notnull(value):
                        ax2.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), f"{value:.1f}",
                                 ha='center', va='bottom', fontsize=8)
                esai_stream = BytesIO()
                plt.tight_layout()
                plt.savefig(esai_stream, format='png')
                plt.close(fig2)
                esai_stream.seek(0)
                doc.add_paragraph("\n📈 Grafik Skor Esai:")
                doc.add_picture(esai_stream, width=Inches(6))

        # ======= Bagian Isian =======
        if not df_isian.empty:
            doc.add_page_break()
            doc.add_heading("✍️ Analisis Isian", level=1)
            table_isian = doc.add_table(rows=1, cols=len(df_isian.columns))
            table_isian.style = 'Table Grid'
            for i, col in enumerate(df_isian.columns):
                table_isian.rows[0].cells[i].text = col
            for row in df_isian.itertuples(index=False):
                cells = table_isian.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)

            # Grafik Isian jika skor numerik tersedia
            if skor_isian and any(pd.notnull(skor_isian)):
                fig3, ax3 = plt.subplots(figsize=(6, 4))
                bars3 = ax3.bar(df_isian["Soal"], skor_isian, color="lightgreen")
                ax3.set_title("Grafik Skor Isian")
                ax3.set_xlabel("Soal Isian")
                ax3.set_ylabel("Skor (%)")
                plt.xticks(rotation=45)
                for bar, value in zip(bars3, skor_isian):
                    if pd.notnull(value):
                        ax3.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), f"{value:.1f}",
                                 ha='center', va='bottom', fontsize=8)
                isian_stream = BytesIO()
                plt.tight_layout()
                plt.savefig(isian_stream, format='png')
                plt.close(fig3)
                doc.add_paragraph("\n📈 Grafik Skor Isian:")
                doc.add_picture(isian_stream, width=Inches(6))

        # Simpan dokumen
        word_stream = BytesIO()
        doc.save(word_stream)
        word_stream.seek(0)

        st.download_button(
            "📥 Unduh sebagai Word",
            data=word_stream.getvalue(),
            file_name="hasil_analisis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Tidak ada hasil yang tersedia untuk diunduh.")

# -------------------------
# Halaman Tentang
# -------------------------
elif menu == "ℹ️ Tentang":
    st.subheader("ℹ️ Tentang Aplikasi")
    st.markdown("""
        Aplikasi ini bertujuan membantu guru dan pendidik dalam menganalisis kualitas soal ujian secara otomatis dan interaktif.

        **Dikembangkan oleh**: [Abu Aisy]  
        **Teknologi**: Python, Streamlit, Plotly  
        **Versi**: 1.1
    """)
